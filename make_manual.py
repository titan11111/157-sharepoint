from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

OUTPUT = '/Users/gosho/Desktop/GitHub-game/157-sharepoint/sharepoint-manual.docx'

GOLD    = RGBColor(0xC7, 0x9A, 0x42)
NAVY    = RGBColor(0x1F, 0x4E, 0x79)
DARK    = RGBColor(0x28, 0x28, 0x28)
GRAY    = RGBColor(0x66, 0x66, 0x66)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── ページ設定（A4）──
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── ヘッダー ──
header = section.header
hp = header.paragraphs[0]
hp.text = '法人営業ナレッジポータル 実装マニュアル'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = hp.runs[0]
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = 'Meiryo'

# ── ページ番号（フッター）──
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
frun = fp.add_run()
frun.font.size = Pt(9)
frun.font.color.rgb = GRAY
frun.font.name = 'Meiryo'
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
frun._r.append(fldChar1)
frun._r.append(instrText)
frun._r.append(fldChar2)

# ─── ヘルパー ───────────────────────────────

def meiryo(run, size=11, bold=False, color=None, italic=False):
    run.font.name = 'Meiryo'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    meiryo(run, size=16, bold=True, color=GOLD)
    # 下線（金色）
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'C79A42')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('▌ ' + text)
    meiryo(run, size=12, bold=True, color=NAVY)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    meiryo(run, size=10.5, color=DARK)
    return p

def add_step(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f'  {num}.  ')
    meiryo(r1, size=10.5, bold=True, color=GOLD)
    r2 = p.add_run(text)
    meiryo(r2, size=10.5, color=DARK)
    return p

def add_point(doc, text, label='POINT'):
    """グレー背景の囲み"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5EAD6')  # gold-light
    p._p.get_or_add_pPr().append(shd)
    r1 = p.add_run(f'【{label}】 ')
    meiryo(r1, size=10, bold=True, color=GOLD)
    r2 = p.add_run(text)
    meiryo(r2, size=10, color=DARK)
    return p

def add_chapter_intro(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0EFED')
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run('📌 この章でやること：' + text)
    meiryo(r, size=10, italic=True, color=DARK)

def add_table_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        meiryo(run, size=10)
    return row

def blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

# ─── 表紙 ─────────────────────────────────────

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('法人営業ナレッジポータル')
meiryo(r, size=22, bold=True, color=GOLD)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run('SharePoint Online 実装マニュアル')
meiryo(r2, size=18, bold=True, color=NAVY)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(40)
r3 = p3.add_run('〜 新卒でもできる！ステップバイステップガイド 〜')
meiryo(r3, size=12, color=GRAY, italic=True)

# 表紙情報表
tbl = doc.add_table(rows=4, cols=2)
tbl.style = 'Table Grid'
info_rows = [
    ('対象者', '入社1〜3年目、SharePoint未経験者'),
    ('想定作業時間', '合計 約4〜5時間'),
    ('使用ツール', 'SharePoint Online / Microsoft Stream / Microsoft Forms / Microsoft Lists'),
    ('作成日', date.today().strftime('%Y年%m月%d日')),
]
for i, (k, v) in enumerate(info_rows):
    r = tbl.rows[i]
    r.cells[0].paragraphs[0].add_run(k).font.bold = True
    for run in r.cells[0].paragraphs[0].runs:
        meiryo(run, size=10, bold=True, color=NAVY)
    r.cells[1].paragraphs[0].add_run(v)
    for run in r.cells[1].paragraphs[0].runs:
        meiryo(run, size=10)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5EAD6')
    r.cells[0]._tc.get_or_add_tcPr().append(shd)

doc.add_page_break()

# ─── 目次（手動）─────────────────────────────

add_h1(doc, '目次')
toc_items = [
    ('第1章', '準備編', '30分'),
    ('第2章', 'ナビゲーション設定', '20分'),
    ('第3章', 'ホームページ レイアウト設定', '40分'),
    ('第4章', 'Hero バナー設置', '20分'),
    ('第5章', '研修カード一覧の構築（最重要）', '60分'),
    ('第6章', '動画の埋め込み（Microsoft Stream）', '30分'),
    ('第7章', 'アンケートの設置（Microsoft Forms）', '30分'),
    ('第8章', '右サイドバーの構築', '30分'),
    ('第9章', '公開と権限設定', '20分'),
    ('第10章', '運用・更新方法', '15分'),
]
for ch, title, time in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{ch}　{title}')
    meiryo(r1, size=11, color=DARK)
    r2 = p.add_run(f'  （所要 {time}）')
    meiryo(r2, size=10, color=GRAY)

doc.add_page_break()

# ─── 第1章 ─────────────────────────────────────

add_h1(doc, '第1章　準備編　（所要 30分）')
add_chapter_intro(doc, 'SharePointにログインし、新しいサイトを作成します。')
blank(doc)

add_h2(doc, 'コミュニケーションサイト vs チームサイト　どちらを選ぶ？')
add_body(doc, '今回のポータルには「コミュニケーションサイト」を選びます。')
tbl = doc.add_table(rows=3, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0]
for cell, txt in zip(hdr.cells, ['', 'コミュニケーションサイト', 'チームサイト']):
    run = cell.paragraphs[0].add_run(txt)
    meiryo(run, size=10, bold=True)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5EAD6')
    cell._tc.get_or_add_tcPr().append(shd)
for cells in [
    ('用途', '情報発信・ポータル', 'チーム内作業・ファイル共有'),
    ('今回の選択', '✅ これを選ぶ', '×'),
]:
    row = tbl.add_row()
    for cell, txt in zip(row.cells, cells):
        meiryo(cell.paragraphs[0].add_run(txt), size=10)
blank(doc)

add_h2(doc, 'サイトの新規作成手順')
steps = [
    'ブラウザで Microsoft 365（office.com）にログインする',
    '左メニューから「SharePoint」をクリックする',
    '画面左上の「＋ サイトの作成」をクリックする',
    '「コミュニケーションサイト」を選択する',
    'サイト名に「法人営業ナレッジポータル」と入力する',
    '言語：日本語　／　タイムゾーン：(UTC+09:00) 大阪、札幌、東京 を選ぶ',
    '「完了」をクリックする（サイト生成に1〜2分かかります）',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
add_point(doc, 'サイトURLは後から変更できません。英字で短く設定してください。例：hojin-portal', 'POINT')

doc.add_page_break()

# ─── 第2章 ─────────────────────────────────────

add_h1(doc, '第2章　ナビゲーション設定　（所要 20分）')
add_chapter_intro(doc, '左ナビに研修カテゴリのリンクを追加します。')
blank(doc)

add_h2(doc, 'ナビゲーション項目の追加手順')
steps = [
    'サイトトップページで右上の「編集」ボタンをクリックする',
    '左ナビの下部に表示される「＋ リンクを追加」をクリックする',
    '追加するリンク名とURLを入力し「OK」をクリックする',
    '以下の順番で10項目を追加する（表参照）',
    '完了したら「保存」をクリックする',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
blank(doc)

add_body(doc, '追加するナビゲーション一覧：')
tbl = doc.add_table(rows=11, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0]
for cell, txt in zip(hdr.cells, ['順番', '表示名', 'リンク先']):
    meiryo(cell.paragraphs[0].add_run(txt), size=10, bold=True)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5EAD6')
    cell._tc.get_or_add_tcPr().append(shd)
nav_items = [
    ('1', 'ホーム', 'サイトトップページ'),
    ('2', '今月の必須研修', '必須研修ページのURL'),
    ('3', '提案力強化', 'カテゴリページのURL'),
    ('4', '医療法人攻略', 'カテゴリページのURL'),
    ('5', '相続・事業承継', 'カテゴリページのURL'),
    ('6', 'コンプライアンス', 'カテゴリページのURL'),
    ('7', '管理職向け', 'カテゴリページのURL'),
    ('8', '全研修一覧', '研修一覧ページのURL'),
    ('9', 'FAQ', 'FAQページのURL'),
    ('10', 'お問い合わせ', 'mailto:training@company.co.jp'),
]
for num, name, url in nav_items:
    row = tbl.add_row()
    for cell, txt in zip(row.cells, [num, name, url]):
        meiryo(cell.paragraphs[0].add_run(txt), size=10)

add_point(doc, '最初はURLが未確定でも、仮に「#」を入れておき後で変更できます。', 'POINT')

doc.add_page_break()

# ─── 第3章 ─────────────────────────────────────

add_h1(doc, '第3章　ホームページ レイアウト設定　（所要 40分）')
add_chapter_intro(doc, 'ページを「左ナビ＋メインコンテンツ＋右サイドバー」の3カラムに設定します。')
blank(doc)

add_h2(doc, 'セクションの追加（3カラム構成）')
steps = [
    'ページ右上の「編集」をクリックしてページ編集モードに入る',
    'ページ中央の「＋」ボタンをクリックして新しいセクションを追加する',
    '「3列」レイアウトを選択する',
    '列の比率設定で「左：中央：右 ＝ 1：3：1」を選ぶ',
    '必要なセクション数だけ同じ手順を繰り返す',
    '最後に「保存してページを閉じる」をクリックする',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)

blank(doc)
add_h2(doc, 'ページ背景の調整')
add_body(doc, 'SharePoint標準のテーマで背景色を白系に設定します。')
steps = [
    '画面右上の歯車アイコン（設定）→「サイトのデザインを変更する」をクリック',
    '「テーマ」から「白」または「ニュートラル」を選択する',
    '変更後、ページをリロードして確認する',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
add_point(doc, '背景色の細かい調整（#fcfbf8など）はSharePoint標準機能では難しいため、最も近い「白」テーマを選択してください。', '注意')

doc.add_page_break()

# ─── 第4章 ─────────────────────────────────────

add_h1(doc, '第4章　Hero バナー設置　（所要 20分）')
add_chapter_intro(doc, 'トップページに大きな研修バナーを設置します。スライド形式で3枚切り替えできます。')
blank(doc)

add_h2(doc, 'Hero ウェブパーツの追加手順')
steps = [
    'ページ編集モードで「＋ ウェブパーツの追加」をクリックする',
    '検索ボックスに「Hero」と入力し「Hero」を選択する',
    'レイアウト選択で「1列（フルワイド）」を選ぶ',
    '「リンクを選択」→「画像をアップロード」で研修のメイン画像を設定する',
    '以下の項目を設定する：\n　　見出し：「今月の重点研修」\n　　説明文：研修内容の概要\n　　ボタンテキスト：「今すぐ視聴する」\n　　ボタンリンク：動画のURL',
    '2枚目・3枚目のスライドも同様に設定する（最大5枚まで）',
    '「保存」をクリックする',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
add_point(doc, '画像は横長（16:9）のものを推奨。推奨サイズは 1920×1080px。高解像度の写真を使うと見栄えが上がります。', 'POINT')

doc.add_page_break()

# ─── 第5章 ─────────────────────────────────────

add_h1(doc, '第5章　研修カード一覧の構築　（所要 60分・最重要）')
add_chapter_intro(doc, '研修18本をカード形式で一覧表示します。Microsoft Lists でデータを管理し、SharePointに連携します。')
blank(doc)

add_h2(doc, 'Step 1：Microsoft Lists で研修リストを作成する')
steps = [
    'SharePointサイト右上の歯車 →「サイトのコンテンツ」をクリック',
    '「新規」→「リスト」→「空のリスト」を選択',
    'リスト名を「研修リスト」と入力し「作成」をクリック',
    '以下の列を追加する（表参照）',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
blank(doc)

add_body(doc, '追加する列の定義：')
tbl = doc.add_table(rows=8, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0]
for cell, txt in zip(hdr.cells, ['列名', 'データ型', '内容・例']):
    meiryo(cell.paragraphs[0].add_run(txt), size=10, bold=True)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5EAD6')
    cell._tc.get_or_add_tcPr().append(shd)
col_defs = [
    ('タイトル', 'テキスト（既存）', '法人営業の基礎と心構え'),
    ('カテゴリ', '選択肢', '提案力強化 / 医療法人 / 相続 / コンプライアンス'),
    ('所要時間', '数値', '15（分単位で入力）'),
    ('動画URL', 'ハイパーリンク', 'Microsoft StreamのビデオURL'),
    ('フォームURL', 'ハイパーリンク', 'Microsoft FormsのフォームURL'),
    ('必須フラグ', 'はい/いいえ', '必須研修の場合「はい」を選択'),
    ('期限', '日付', '2026/05/31'),
]
for name, dtype, example in col_defs:
    row = tbl.add_row()
    for cell, txt in zip(row.cells, [name, dtype, example]):
        meiryo(cell.paragraphs[0].add_run(txt), size=10)

blank(doc)
add_h2(doc, 'Step 2：18件のデータを入力する')
steps = [
    '作成したリストを開き「新しい項目」をクリック',
    '研修1本ずつタイトル・カテゴリ・所要時間・URL・必須フラグ・期限を入力',
    '18本すべて入力したら完了',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
add_point(doc, '動画URLとフォームURLは、第6章・第7章で取得します。先にリストの骨格だけ作っておき、後からURLを入力してもOKです。', 'POINT')

blank(doc)
add_h2(doc, 'Step 3：ギャラリービューを設定する')
steps = [
    'リスト右上の「＋ ビューを追加」をクリック',
    '「ギャラリー」を選択する',
    'ビュー名を「カード表示」と入力',
    '表示する列：タイトル / カテゴリ / 所要時間 / 必須フラグ / 期限 を選択',
    '「保存」をクリック',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)

blank(doc)
add_h2(doc, 'Step 4：ホームページにリストを埋め込む')
steps = [
    'ホームページを編集モードで開く',
    '研修一覧を表示したいエリアで「＋ ウェブパーツの追加」をクリック',
    '「リスト」ウェブパーツを選択する',
    '先ほど作成した「研修リスト」を選択',
    'ビューを「カード表示（ギャラリー）」に変更する',
    '表示件数・レイアウトを調整して「保存」をクリック',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
add_point(doc, '「動画を見る」「アンケート回答」ボタンは、ギャラリービューのカード内リンクとして表示されます。URLが正しく入力されていることを確認してください。', '確認')

doc.add_page_break()

# ─── 第6章 ─────────────────────────────────────

add_h1(doc, '第6章　動画の埋め込み（Microsoft Stream）　（所要 30分）')
add_chapter_intro(doc, '研修動画をMicrosoft Streamにアップロードし、URLを取得して研修リストに登録します。')
blank(doc)

add_h2(doc, 'Microsoft Stream へのアップロード手順')
steps = [
    'ブラウザで stream.microsoft.com にアクセスしてログイン',
    '「マイコンテンツ」→「ビデオ」→「新しいビデオ」→「ファイルのアップロード」をクリック',
    '動画ファイル（MP4推奨）を選択してアップロード開始',
    'アップロード完了後、動画の詳細ページを開く',
    '「共有」→「リンクのコピー」でURLを取得する',
    '取得したURLを研修リストの「動画URL」列に貼り付ける',
    '18本すべて同様に繰り返す',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
add_point(doc, 'アップロードできるファイルサイズは最大100GBです。動画は圧縮してからアップロードすると処理が速くなります。', 'POINT')
add_point(doc, '動画の共有設定で「組織内の全員が視聴可能」に設定しないと、受講者が再生できません。必ず確認してください。', '重要')

doc.add_page_break()

# ─── 第7章 ─────────────────────────────────────

add_h1(doc, '第7章　アンケートの設置（Microsoft Forms）　（所要 30分）')
add_chapter_intro(doc, 'Microsoft Formsでアンケートを作成し、各研修カードにリンクします。')
blank(doc)

add_h2(doc, 'Forms でアンケートを新規作成する手順')
steps = [
    'forms.office.com にアクセスしてログイン',
    '「新しいフォーム」をクリックする',
    'フォーム名を「研修01_法人営業の基礎と心構え_アンケート」のように入力',
    '質問を追加する（例：感想・理解度・実践意欲など3〜5問を目安に）',
    '「共有」ボタンをクリック',
    '「リンクのコピー」でフォームURLを取得する',
    '取得したURLを研修リストの「フォームURL」列に貼り付ける',
    '18本分すべて同様に作成する',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)

blank(doc)
add_h2(doc, '回答データの確認方法')
steps = [
    'forms.office.com でフォームを開く',
    '「回答」タブをクリックする',
    '個人別・質問別の集計がグラフで表示される',
    '「Excelで開く」でExcelに書き出して分析することも可能',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
add_point(doc, 'アンケート結果の共有は「Forms の共同編集者」として担当者を追加することで複数名で確認できます。', 'POINT')

doc.add_page_break()

# ─── 第8章 ─────────────────────────────────────

add_h1(doc, '第8章　右サイドバーの構築　（所要 30分）')
add_chapter_intro(doc, '右サイドバーに「お知らせ」「人気研修TOP3」「よくある質問」を設置します。')
blank(doc)

add_h2(doc, 'お知らせ（Newsウェブパーツ）の設置')
steps = [
    'ホームページ編集モードで右列エリアの「＋」をクリック',
    '「ニュース」ウェブパーツを追加する',
    'レイアウトを「リスト」に変更する（コンパクト表示）',
    '表示件数を「4件」に設定する',
    '「保存」をクリックする',
    'お知らせを投稿するには：「＋新規」→「ニュース投稿」から追加',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)

blank(doc)
add_h2(doc, '人気研修TOP3（Quick Linksウェブパーツ）の設置')
steps = [
    '右列エリアで「＋ ウェブパーツの追加」→「クイックリンク」を選択',
    'レイアウト「リスト」を選択する',
    '「＋ リンクの追加」で研修動画のURLを3件追加する',
    '各リンクのタイトルを「① 医療法人開拓の最新事例」のように設定',
    '順番はドラッグ&ドロップで変更可能',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)

blank(doc)
add_h2(doc, 'よくある質問（Highlighted Contentウェブパーツ）の設置')
steps = [
    '右列エリアで「＋ ウェブパーツの追加」→「強調表示されたコンテンツ」を選択',
    'コンテンツソースを「このサイト」に設定',
    'フィルターでFAQページを絞り込む',
    '表示件数を3〜5件に設定して「保存」',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)

doc.add_page_break()

# ─── 第9章 ─────────────────────────────────────

add_h1(doc, '第9章　公開と権限設定　（所要 20分）')
add_chapter_intro(doc, 'サイトを全社員が閲覧できるよう公開し、権限を設定します。')
blank(doc)

add_h2(doc, 'サイトの公開手順')
steps = [
    'ページ右上の「発行」ボタンをクリックする',
    '「このページを発行しますか？」の確認メッセージが出たら「発行」をクリック',
    '公開後はURLを共有すれば誰でもアクセスできます',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)
add_point(doc, '各ページを作成・編集するたびに「発行」が必要です。下書き状態のページは自分しか見えません。', '注意')

blank(doc)
add_h2(doc, '権限の設定')
steps = [
    '歯車アイコン →「サイトのアクセス許可」をクリック',
    '「訪問者を招待する」→「組織全体」を追加する（全社員が閲覧可能になる）',
    '編集者は「メンバー」として個別に追加する',
    'Microsoft Lists・Stream・Formsも同様に権限を確認する',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)

doc.add_page_break()

# ─── 第10章 ─────────────────────────────────────

add_h1(doc, '第10章　運用・更新方法　（所要 15分）')
add_chapter_intro(doc, '研修を新規追加する方法・お知らせ投稿・よくあるトラブル対応を説明します。')
blank(doc)

add_h2(doc, '研修を新規追加する手順')
steps = [
    'Microsoft Stream に動画をアップロードしてURLを取得',
    'Microsoft Forms でアンケートを作成してURLを取得',
    'SharePointの「研修リスト」を開き「新しい項目」をクリック',
    '各列（タイトル・カテゴリ・URL等）を入力して「保存」',
    'ホームページに自動反映されることを確認する',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)

blank(doc)
add_h2(doc, 'お知らせを投稿する手順')
steps = [
    'ホームページの「ニュース」ウェブパーツ横の「＋すべて表示」をクリック',
    '「＋新規」→「ニュース投稿」を選択',
    'タイトル・本文・画像を入力して「発行」をクリック',
]
for i, s in enumerate(steps, 1):
    add_step(doc, i, s)

blank(doc)
add_h2(doc, 'よくあるトラブルと対処法')
tbl = doc.add_table(rows=6, cols=2)
tbl.style = 'Table Grid'
hdr = tbl.rows[0]
for cell, txt in zip(hdr.cells, ['症状', '対処法']):
    meiryo(cell.paragraphs[0].add_run(txt), size=10, bold=True)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5EAD6')
    cell._tc.get_or_add_tcPr().append(shd)
troubles = [
    ('動画が再生できない', 'Streamの共有設定を確認。「組織内の全員」が再生可能か確認してください'),
    ('カードが表示されない', 'リストのビュー設定を確認。ギャラリービューが選択されているか確認'),
    ('ナビにリンクが出ない', 'ページが「発行」されているか確認。下書き状態は表示されません'),
    ('アンケートが開かない', 'FormsのURLが正しくコピーされているか確認。短縮URLは使用不可'),
    ('権限エラーが出る', 'サイトのアクセス許可で閲覧者が追加されているか確認してください'),
]
for symptom, solution in troubles:
    row = tbl.add_row()
    meiryo(row.cells[0].paragraphs[0].add_run(symptom), size=10)
    meiryo(row.cells[1].paragraphs[0].add_run(solution), size=10)

blank(doc)
add_point(doc, '困ったときは training@company.co.jp または 03-1234-5678 に問い合わせてください。', 'サポート')

blank(doc)
blank(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('以上　実装マニュアル 完')
meiryo(r, size=11, color=GRAY, italic=True)

# ─── 保存 ─────────────────────────────────────
doc.save(OUTPUT)
print(f'✅ 保存完了: {OUTPUT}')
